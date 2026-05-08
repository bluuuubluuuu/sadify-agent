from __future__ import annotations

from dataclasses import dataclass
from datetime import UTC, datetime
from html import escape
from io import BytesIO
from pathlib import Path
from textwrap import wrap
from typing import Sequence
import re

from docx import Document

from sadify.renderers.wiki_markdown import WikiNoteDraft
from sadify.schemas import ExportRecord, SadVersionRecord


class ExportGenerationError(ValueError):
    pass


@dataclass(frozen=True)
class PreparedExportArtifact:
    export_id: str
    export_type: str
    relative_path: str
    file_name: str
    mime_type: str
    content: bytes


@dataclass(frozen=True)
class ExportPackage:
    artifacts: tuple[PreparedExportArtifact, ...]
    records: tuple[ExportRecord, ...]

    def export_types(self) -> tuple[str, ...]:
        return tuple(artifact.export_type for artifact in self.artifacts)

    def artifact_by_type(self, export_type: str) -> PreparedExportArtifact:
        for artifact in self.artifacts:
            if artifact.export_type == export_type:
                return artifact
        raise ExportGenerationError(f"Export artifact not found: {export_type}")

    def artifact_by_relative_path(self, relative_path: str) -> PreparedExportArtifact:
        normalized_path = _normalize_relative_path(relative_path)
        for artifact in self.artifacts:
            if artifact.relative_path == normalized_path:
                return artifact
        raise ExportGenerationError(f"Export artifact not found: {relative_path}")


def prepare_export_package(
    *,
    sad_version: SadVersionRecord,
    wiki_notes: Sequence[WikiNoteDraft] = (),
    project_slug: str,
    created_at: datetime | None = None,
    created_by: str,
    source_knowledge_item_version_ids: Sequence[str] = (),
    first_export_number: int = 1,
) -> ExportPackage:
    timestamp = created_at or datetime.now(UTC)
    safe_slug = _safe_slug(project_slug) or _safe_slug(sad_version.sad_version_id)
    base_name = f"SAD-v{sad_version.version_number}-{safe_slug}"

    prepared_specs = [
        (
            "google_doc",
            f"{base_name}.google-doc.html",
            "sad",
            "text/html",
            _render_google_doc_html(sad_version),
        ),
        (
            "pdf",
            f"{base_name}.pdf",
            "sad",
            "application/pdf",
            _render_pdf(sad_version.rendered_markdown),
        ),
        (
            "docx",
            f"{base_name}.docx",
            "sad",
            (
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            _render_docx(sad_version.rendered_markdown),
        ),
    ]
    prepared_specs.extend(
        (
            "wiki_markdown",
            note.file_name,
            f"wiki/{note.folder}",
            "text/markdown",
            note.markdown.encode("utf-8"),
        )
        for note in wiki_notes
    )

    artifacts: list[PreparedExportArtifact] = []
    records: list[ExportRecord] = []
    for offset, spec in enumerate(prepared_specs):
        export_type, file_name, folder, mime_type, content = spec
        export_id = f"EXP-{first_export_number + offset:03d}"
        relative_path = _normalize_relative_path(f"{folder}/{file_name}")
        artifact = PreparedExportArtifact(
            export_id=export_id,
            export_type=export_type,
            relative_path=relative_path,
            file_name=file_name,
            mime_type=mime_type,
            content=content,
        )
        record = ExportRecord(
            export_id=export_id,
            export_type=export_type,
            source_sad_version_id=sad_version.sad_version_id,
            source_knowledge_item_version_ids=list(source_knowledge_item_version_ids),
            file_name=file_name,
            drive_file_id=None,
            url=None,
            created_at=timestamp,
            created_by=created_by,
            status="success",
            error_message=None,
        )
        artifacts.append(artifact)
        records.append(record)

    return ExportPackage(artifacts=tuple(artifacts), records=tuple(records))


def write_export_package(
    package: ExportPackage,
    output_dir: str | Path,
) -> tuple[Path, ...]:
    root = Path(output_dir).resolve()
    root.mkdir(parents=True, exist_ok=True)

    written_paths: list[Path] = []
    for artifact in package.artifacts:
        target_path = (root / artifact.relative_path).resolve()
        _ensure_inside_root(target_path, root)
        target_path.parent.mkdir(parents=True, exist_ok=True)
        target_path.write_bytes(artifact.content)
        written_paths.append(target_path)
    return tuple(written_paths)


def _render_google_doc_html(sad_version: SadVersionRecord) -> bytes:
    lines = [
        "<!doctype html>",
        '<html lang="en">',
        "<head>",
        '<meta charset="utf-8">',
        f"<title>{escape(_title_from_markdown(sad_version.rendered_markdown))}</title>",
        "</head>",
        "<body>",
    ]

    in_list = False
    for line in sad_version.rendered_markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            if in_list:
                lines.append("</ul>")
                in_list = False
            continue

        heading = _heading_level_and_text(stripped)
        if heading is not None:
            if in_list:
                lines.append("</ul>")
                in_list = False
            level, text = heading
            lines.append(f"<h{level}>{escape(text)}</h{level}>")
            continue

        if stripped.startswith("- "):
            if not in_list:
                lines.append("<ul>")
                in_list = True
            lines.append(f"<li>{escape(stripped[2:])}</li>")
            continue

        if in_list:
            lines.append("</ul>")
            in_list = False
        lines.append(f"<p>{escape(stripped)}</p>")

    if in_list:
        lines.append("</ul>")
    lines.extend(["</body>", "</html>"])
    return "\n".join(lines).encode("utf-8")


def _render_docx(markdown: str) -> bytes:
    document = Document()
    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        heading = _heading_level_and_text(stripped)
        if heading is not None:
            level, text = heading
            document.add_heading(text, level=min(level, 4))
            continue

        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
            continue

        document.add_paragraph(stripped)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_pdf(markdown: str) -> bytes:
    text_lines = _plain_text_lines(markdown)
    stream_lines = ["BT", "/F1 10 Tf", "14 TL", "72 720 Td"]
    for index, line in enumerate(text_lines[:45]):
        if index == 0:
            stream_lines.append(f"({_pdf_text(line)}) Tj")
        else:
            stream_lines.append(f"T* ({_pdf_text(line)}) Tj")
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("latin-1", errors="replace")

    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        (
            f"<< /Length {len(stream)} >>\nstream\n".encode("ascii")
            + stream
            + b"\nendstream"
        ),
    ]
    return _build_pdf(objects)


def _build_pdf(objects: Sequence[bytes]) -> bytes:
    output = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []

    for index, body in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode("ascii"))
        output.extend(body)
        output.extend(b"\nendobj\n")

    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(output)


def _plain_text_lines(markdown: str) -> list[str]:
    lines: list[str] = []
    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        stripped = stripped.lstrip("#").strip()
        stripped = stripped.removeprefix("- ").strip()
        lines.extend(wrap(stripped, width=88) or [""])
    return lines


def _heading_level_and_text(line: str) -> tuple[int, str] | None:
    match = re.match(r"^(#{1,6})\s+(.+)$", line)
    if match is None:
        return None
    return len(match.group(1)), match.group(2)


def _title_from_markdown(markdown: str) -> str:
    for line in markdown.splitlines():
        heading = _heading_level_and_text(line.strip())
        if heading is not None and heading[0] == 1:
            return heading[1]
    return "SAD Export"


def _pdf_text(value: str) -> str:
    normalized = value.encode("latin-1", errors="replace").decode("latin-1")
    return (
        normalized
        .replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
    )


def _safe_slug(value: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")
    return slug or "export"


def _normalize_relative_path(value: str) -> str:
    return value.replace("\\", "/").strip("/")


def _ensure_inside_root(target_path: Path, root: Path) -> None:
    try:
        target_path.relative_to(root)
    except ValueError as exc:
        raise ExportGenerationError(
            f"Export path escapes target directory: {target_path}"
        ) from exc
