from __future__ import annotations

from dataclasses import dataclass
import csv
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any


SUPPORTED_FILE_TYPES = ("MD", "TXT", "PDF", "DOCX", "XLSX", "CSV")


class FileExtractionError(ValueError):
    """Raised when a supported file cannot produce readable text."""


class UnsupportedFileTypeError(FileExtractionError):
    """Raised when a file extension is outside the MVP scope."""


@dataclass(frozen=True)
class ExtractedRequirementSource:
    filename: str
    file_type: str
    normalized_text: str
    metadata: dict[str, Any]

    def to_display_dict(self) -> dict[str, Any]:
        return {
            "filename": self.filename,
            "file_type": self.file_type,
            "normalized_text": self.normalized_text,
            "metadata": self.metadata,
        }


def extract_requirement_source(
    filename: str,
    content: bytes,
) -> ExtractedRequirementSource:
    clean_filename = filename.strip() or "uploaded file"
    extension = Path(clean_filename).suffix.lower()
    metadata: dict[str, Any] = {
        "filename": clean_filename,
        "byte_count": len(content),
    }

    try:
        file_type, text, extracted_metadata = _extract_by_extension(
            clean_filename,
            extension,
            content,
        )
    except FileExtractionError:
        raise
    except Exception as exc:  # pragma: no cover - defensive wrapper
        raise FileExtractionError(
            f"Could not extract readable text from {clean_filename}: {exc}"
        ) from exc

    normalized_text = _normalize_text(text)
    if not normalized_text:
        if file_type == "pdf":
            raise FileExtractionError(_pdf_text_failure_message(clean_filename))
        raise FileExtractionError(
            f"No readable requirement text found in {clean_filename}."
        )

    metadata.update(extracted_metadata)
    return ExtractedRequirementSource(
        filename=clean_filename,
        file_type=file_type,
        normalized_text=normalized_text,
        metadata=metadata,
    )


def _extract_by_extension(
    filename: str,
    extension: str,
    content: bytes,
) -> tuple[str, str, dict[str, Any]]:
    if extension in {".md", ".markdown"}:
        text = _decode_text(content)
        return "md", text, {"character_count": len(text)}
    if extension == ".txt":
        text = _decode_text(content)
        return "txt", text, {"character_count": len(text)}
    if extension == ".pdf":
        return "pdf", *_extract_pdf(filename, content)
    if extension == ".docx":
        return "docx", *_extract_docx(content)
    if extension == ".xlsx":
        return "xlsx", *_extract_xlsx(content)
    if extension == ".csv":
        return "csv", *_extract_csv(content)

    display_extension = extension or "<none>"
    raise UnsupportedFileTypeError(
        f"Unsupported file type '{display_extension}' for {filename}. "
        f"Supported files: {', '.join(SUPPORTED_FILE_TYPES)}."
    )


def _decode_text(content: bytes) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError:
        return content.decode("latin-1")


def _extract_pdf(filename: str, content: bytes) -> tuple[str, dict[str, Any]]:
    from pypdf import PdfReader

    try:
        reader = PdfReader(BytesIO(content), strict=False)
        page_count = len(reader.pages)
        page_text: list[str] = []
        failed_pages: list[int] = []
        for page_index in range(page_count):
            try:
                text = reader.pages[page_index].extract_text()
            except Exception:
                failed_pages.append(page_index + 1)
                continue
            if text and text.strip():
                page_text.append(text.strip())
        return "\n\n".join(page_text), {
            "page_count": page_count,
            "failed_pages": failed_pages,
        }
    except Exception as exc:
        raise FileExtractionError(_pdf_text_failure_message(filename)) from exc


def _pdf_text_failure_message(filename: str) -> str:
    return (
        f"Could not read selectable text from {filename}. "
        "The PDF may be scanned, protected, damaged, or exported in a format "
        "this prototype cannot parse yet. Try uploading a DOCX, MD, TXT, CSV, "
        "or XLSX version, or export the PDF again with selectable text."
    )


def _extract_docx(content: bytes) -> tuple[str, dict[str, Any]]:
    from docx import Document

    document = Document(BytesIO(content))
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    table_summaries: list[str] = []
    table_row_count = 0

    for table_index, table in enumerate(document.tables, start=1):
        table_rows: list[list[str]] = []
        for row in table.rows:
            row_values = [
                _normalize_cell_text(cell.text)
                for cell in row.cells
                if _normalize_cell_text(cell.text)
            ]
            if row_values:
                table_rows.append(row_values)
        if table_rows:
            table_row_count += len(table_rows)
            table_summaries.append(
                "\n".join(
                    [f"DOCX table {table_index}"]
                    + [", ".join(row) for row in table_rows]
                )
            )

    text_parts = ["\n".join(paragraphs)] if paragraphs else []
    text_parts.extend(table_summaries)
    return "\n\n".join(text_parts), {
        "paragraph_count": len(paragraphs),
        "table_count": len(document.tables),
        "table_row_count": table_row_count,
    }


def _extract_csv(content: bytes) -> tuple[str, dict[str, Any]]:
    text = _decode_text(content)
    rows = [
        [cell.strip() for cell in row]
        for row in csv.reader(StringIO(text))
        if any(cell.strip() for cell in row)
    ]
    if not rows:
        return "", {"row_count": 0, "columns": []}

    columns = rows[0]
    data_rows = rows[1:]
    summary = _table_summary_text(
        title="CSV file",
        columns=columns,
        data_rows=data_rows,
    )
    return summary, {
        "row_count": len(data_rows),
        "columns": columns,
    }


def _extract_xlsx(content: bytes) -> tuple[str, dict[str, Any]]:
    from openpyxl import load_workbook

    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    sheet_summaries: list[str] = []
    sheet_metadata: list[dict[str, Any]] = []

    for sheet in workbook.worksheets:
        rows = [
            [_stringify_cell(cell) for cell in row]
            for row in sheet.iter_rows(values_only=True)
        ]
        rows = [row for row in rows if any(cell for cell in row)]
        if not rows:
            sheet_metadata.append(
                {"name": sheet.title, "row_count": 0, "columns": []}
            )
            continue

        columns = rows[0]
        data_rows = rows[1:]
        sheet_summaries.append(
            _table_summary_text(
                title=f"Sheet: {sheet.title}",
                columns=columns,
                data_rows=data_rows,
            )
        )
        sheet_metadata.append(
            {
                "name": sheet.title,
                "row_count": len(data_rows),
                "columns": columns,
            }
        )

    return "\n\n".join(sheet_summaries), {
        "sheet_count": len(workbook.worksheets),
        "sheets": sheet_metadata,
    }


def _table_summary_text(
    *,
    title: str,
    columns: list[str],
    data_rows: list[list[str]],
    preview_limit: int = 5,
) -> str:
    lines = [
        title,
        f"Rows: {len(data_rows)}",
        f"Columns: {', '.join(columns)}",
    ]
    preview_rows = data_rows[:preview_limit]
    if preview_rows:
        lines.append("Preview:")
        lines.extend(", ".join(row) for row in preview_rows)
    return "\n".join(lines)


def _stringify_cell(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _normalize_cell_text(value: str) -> str:
    return " ".join(value.split())


def _normalize_text(text: str) -> str:
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").split("\n")]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    return "\n".join(lines)
