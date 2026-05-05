from io import BytesIO

import pytest

from sadify.extractors.business_files import (
    FileExtractionError,
    UnsupportedFileTypeError,
    extract_requirement_source,
)


def test_extracts_plain_text_and_markdown_files():
    text_source = extract_requirement_source(
        "warehouse.txt",
        b"Warehouse staff need to record stock movement.",
    )
    markdown_source = extract_requirement_source(
        "requirement.md",
        b"# Plantation App\n\nField supervisors need daily harvest records.",
    )

    assert text_source.file_type == "txt"
    assert text_source.normalized_text == (
        "Warehouse staff need to record stock movement."
    )
    assert text_source.metadata["character_count"] == 46
    assert markdown_source.file_type == "md"
    assert "Field supervisors need daily harvest records." in (
        markdown_source.normalized_text
    )


def test_extracts_selectable_pdf_text():
    pdf_source = extract_requirement_source(
        "stock.pdf",
        _sample_pdf_bytes(
            "Warehouse stock movement needs supervisor approval."
        ),
    )

    assert pdf_source.file_type == "pdf"
    assert "Warehouse stock movement" in pdf_source.normalized_text
    assert pdf_source.metadata["page_count"] == 1


def test_rejects_unreadable_pdf_with_business_friendly_message():
    with pytest.raises(FileExtractionError) as exc_info:
        extract_requirement_source(
            "Project_Statement_of_HAMS_Wialon_Report.pdf",
            b"%PDF-1.4\n1 0 obj\n<< /Length 200 >>\nstream\nBT\n",
        )

    message = str(exc_info.value)
    assert message == (
        "Could not read selectable text from "
        "Project_Statement_of_HAMS_Wialon_Report.pdf. The PDF may be scanned, "
        "protected, damaged, or exported in a format this prototype cannot parse yet. "
        "Try uploading a DOCX, MD, TXT, CSV, or XLSX version, or export the PDF "
        "again with selectable text."
    )
    assert "Stream has ended unexpectedly" not in message


def test_extracts_docx_paragraphs():
    docx_source = extract_requirement_source(
        "plantation.docx",
        _sample_docx_bytes(
            [
                "Field workers record harvest by block.",
                "Supervisors approve rejected entries.",
            ]
        ),
    )

    assert docx_source.file_type == "docx"
    assert docx_source.normalized_text == (
        "Field workers record harvest by block.\n"
        "Supervisors approve rejected entries."
    )
    assert docx_source.metadata["paragraph_count"] == 2


def test_extracts_docx_table_cells():
    docx_source = extract_requirement_source(
        "plantation-report.docx",
        _sample_docx_with_table_bytes(),
    )

    assert "Field Daily Report" in docx_source.normalized_text
    assert "DOCX table 1" in docx_source.normalized_text
    assert "Block, Activity, Status" in docx_source.normalized_text
    assert "BLI05, Spraying, Completed" in docx_source.normalized_text
    assert "BLI06, Harvesting, Pending" in docx_source.normalized_text
    assert docx_source.metadata["paragraph_count"] == 1
    assert docx_source.metadata["table_count"] == 1
    assert docx_source.metadata["table_row_count"] == 3


def test_extracts_docx_header_and_footer_text():
    docx_source = extract_requirement_source(
        "plantation-report.docx",
        _sample_docx_with_header_footer_bytes(),
    )

    assert "Header: Estate Operations Requirement" in docx_source.normalized_text
    assert "Field workers record rainfall by block." in docx_source.normalized_text
    assert "Footer: Prepared by operations team" in docx_source.normalized_text
    assert docx_source.metadata["header_footer_paragraph_count"] == 2


def test_summarizes_csv_headers_rows_and_preview():
    csv_source = extract_requirement_source(
        "harvest.csv",
        (
            "Block,Worker,WeightKg,Status\n"
            "A01,Aminah,120,Pending\n"
            "A02,Bala,95,Approved\n"
        ).encode("utf-8"),
    )

    assert csv_source.file_type == "csv"
    assert csv_source.metadata["row_count"] == 2
    assert csv_source.metadata["columns"] == [
        "Block",
        "Worker",
        "WeightKg",
        "Status",
    ]
    assert "Rows: 2" in csv_source.normalized_text
    assert "Columns: Block, Worker, WeightKg, Status" in (
        csv_source.normalized_text
    )
    assert "A01, Aminah, 120, Pending" in csv_source.normalized_text


def test_summarizes_csv_quoted_commas_and_multiline_cells():
    csv_source = extract_requirement_source(
        "issues.csv",
        (
            "Block,Issue,Action\n"
            'BLI05,"Loose fruit count, high variance","Supervisor review needed"\n'
            'BLI06,"Line one\nline two","Escalate to manager"\n'
        ).encode("utf-8"),
    )

    assert csv_source.metadata["row_count"] == 2
    assert "BLI05, Loose fruit count, high variance, Supervisor review needed" in (
        csv_source.normalized_text
    )
    assert "BLI06, Line one line two, Escalate to manager" in (
        csv_source.normalized_text
    )


def test_summarizes_xlsx_sheets_headers_rows_and_preview():
    xlsx_source = extract_requirement_source(
        "daily-harvest.xlsx",
        _sample_xlsx_bytes(),
    )

    assert xlsx_source.file_type == "xlsx"
    assert xlsx_source.metadata["sheet_count"] == 1
    assert xlsx_source.metadata["sheets"][0]["name"] == "Harvest"
    assert xlsx_source.metadata["sheets"][0]["row_count"] == 2
    assert xlsx_source.metadata["sheets"][0]["columns"] == [
        "Block",
        "Worker",
        "WeightKg",
    ]
    assert "Sheet: Harvest" in xlsx_source.normalized_text
    assert "A01, Aminah, 120" in xlsx_source.normalized_text


def test_summarizes_xlsx_multiple_sheets_comments_and_formulas():
    xlsx_source = extract_requirement_source(
        "estate-workbook.xlsx",
        _sample_xlsx_with_comments_and_formula_bytes(),
    )

    assert xlsx_source.metadata["sheet_count"] == 2
    assert xlsx_source.metadata["sheets"][1]["name"] == "Issues"
    assert xlsx_source.metadata["sheets"][1]["comment_count"] == 1
    assert "Sheet: Harvest" in xlsx_source.normalized_text
    assert "Sheet: Issues" in xlsx_source.normalized_text
    assert "No mobile signal" in xlsx_source.normalized_text
    assert "Comment B2: Field supervisor says offline mode is needed" in (
        xlsx_source.normalized_text
    )
    assert "=CONCAT(\"Needs \",\"approval\")" in xlsx_source.normalized_text


def test_rejects_unsupported_file_type_with_clear_message():
    with pytest.raises(UnsupportedFileTypeError) as exc_info:
        extract_requirement_source("photo.png", b"not supported")

    assert "Unsupported file type '.png'" in str(exc_info.value)
    assert "MD, TXT, PDF, DOCX, XLSX, CSV" in str(exc_info.value)


def test_rejects_empty_extracted_content_with_clear_message():
    with pytest.raises(FileExtractionError) as exc_info:
        extract_requirement_source("blank.txt", b"   ")

    assert "No readable requirement text found in blank.txt." == str(
        exc_info.value
    )


def _sample_docx_bytes(paragraphs: list[str]) -> bytes:
    from docx import Document

    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _sample_docx_with_table_bytes() -> bytes:
    from docx import Document

    document = Document()
    document.add_paragraph("Field Daily Report")
    table = document.add_table(rows=3, cols=3)
    values = [
        ["Block", "Activity", "Status"],
        ["BLI05", "Spraying", "Completed"],
        ["BLI06", "Harvesting", "Pending"],
    ]
    for row_index, row_values in enumerate(values):
        for cell_index, value in enumerate(row_values):
            table.rows[row_index].cells[cell_index].text = value
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _sample_docx_with_header_footer_bytes() -> bytes:
    from docx import Document

    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "Header: Estate Operations Requirement"
    document.add_paragraph("Field workers record rainfall by block.")
    section.footer.paragraphs[0].text = "Footer: Prepared by operations team"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _sample_xlsx_bytes() -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Harvest"
    sheet.append(["Block", "Worker", "WeightKg"])
    sheet.append(["A01", "Aminah", 120])
    sheet.append(["A02", "Bala", 95])
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _sample_xlsx_with_comments_and_formula_bytes() -> bytes:
    from openpyxl import Workbook
    from openpyxl.comments import Comment

    workbook = Workbook()
    harvest = workbook.active
    harvest.title = "Harvest"
    harvest.append(["Block", "Worker", "WeightKg"])
    harvest.append(["A01", "Aminah", 120])

    issues = workbook.create_sheet("Issues")
    issues.append(["Block", "Issue", "Next Step"])
    issues.append(["BLI05", "No mobile signal", '=CONCAT("Needs ","approval")'])
    issues["B2"].comment = Comment(
        "Field supervisor says offline mode is needed",
        "operations",
    )

    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _sample_pdf_bytes(text: str) -> bytes:
    stream = f"BT\n/F1 12 Tf\n72 720 Td\n({text}) Tj\nET".encode("ascii")
    return (
        b"%PDF-1.4\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n"
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>\n"
        b"endobj\n"
        b"4 0 obj\n"
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\n"
        b"endobj\n"
        b"5 0 obj\n"
        + f"<< /Length {len(stream)} >>\n".encode("ascii")
        + b"stream\n"
        + stream
        + b"\nendstream\nendobj\n"
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"0000000241 00000 n \n"
        b"0000000311 00000 n \n"
        b"trailer\n<< /Root 1 0 R /Size 6 >>\n"
        b"startxref\n443\n%%EOF\n"
    )
